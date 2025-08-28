"""
Converter from Pandoc AST back to DOCX with metadata restoration.

This converter completes the round-trip conversion by:
1. Converting AST back to DOCX using Pandoc
2. Restoring Word-specific metadata and formatting
3. Re-injecting preserved XML fragments
"""

from __future__ import annotations

import json
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional
from xml.etree import ElementTree as ET
import shutil

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..core.document_model import DocumentModel, WordMetadata, XMLFragments


class ASTToDocxConverter:
    """
    Converts our hybrid DocumentModel back to DOCX format.
    
    Uses a multi-stage approach:
    1. Generate base DOCX from AST using Pandoc
    2. Enhance with preserved metadata using python-docx
    3. Re-inject complex XML fragments
    """
    
    def __init__(self, pandoc_path: str = "pandoc"):
        self.pandoc_path = pandoc_path
        self._validate_pandoc()
        
        # OOXML namespaces
        self.namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
        }
    
    def _validate_pandoc(self) -> None:
        """Ensure Pandoc is available."""
        try:
            result = subprocess.run(
                [self.pandoc_path, "--version"], 
                capture_output=True, 
                text=True, 
                check=True
            )
            if "pandoc" not in result.stdout.lower():
                raise RuntimeError("Pandoc validation failed")
        except (subprocess.SubprocessError, FileNotFoundError) as e:
            raise RuntimeError(f"Pandoc not found or not working: {e}")
    
    def convert(self, document_model: DocumentModel, output_path: Path) -> None:
        """
        Convert DocumentModel back to DOCX file.
        
        Args:
            document_model: The document model to convert
            output_path: Path where the DOCX file should be saved
        """
        # Stage 1: Generate base DOCX from AST using Pandoc
        base_docx_path = self._generate_base_docx(document_model)
        
        try:
            # Stage 2: Enhance with metadata and formatting
            enhanced_doc = self._enhance_with_metadata(base_docx_path, document_model.word_metadata)
            
            # Stage 3: Re-inject complex XML fragments  
            final_docx_path = self._inject_xml_fragments(
                enhanced_doc, 
                document_model.xml_fragments,
                output_path
            )
            
            # Ensure the final file is at the requested location
            if final_docx_path != output_path:
                shutil.move(str(final_docx_path), str(output_path))
                
        finally:
            # Cleanup temporary files
            if base_docx_path and base_docx_path.exists():
                base_docx_path.unlink()
    
    def _generate_base_docx(self, document_model: DocumentModel) -> Path:
        """Generate base DOCX from Pandoc AST."""
        # Create temporary file for Pandoc JSON
        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as json_file:
            pandoc_json = document_model.pandoc_ast.to_pandoc_json()
            json.dump(pandoc_json, json_file, indent=2)
            json_path = Path(json_file.name)
        
        # Create temporary output file
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_docx.close()
        temp_docx_path = Path(temp_docx.name)
        
        try:
            # Use Pandoc to convert JSON to DOCX
            cmd = [
                self.pandoc_path,
                str(json_path),
                "--from", "json",
                "--to", "docx",
                "--output", str(temp_docx_path),
                "--standalone",
            ]
            reference_template = self._get_reference_template()
            # Only pass reference-doc if we have a valid template path
            if reference_template:
                try:
                    ref_path = Path(reference_template)
                    if ref_path.exists():
                        cmd.extend(["--reference-doc", str(ref_path)])
                except Exception:
                    # If template resolution fails, proceed without it
                    pass

            result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            
            return temp_docx_path
            
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"Pandoc DOCX generation failed: {e.stderr}")
        finally:
            # Cleanup JSON file
            if json_path.exists():
                json_path.unlink()
    
    def _get_reference_template(self) -> Optional[str]:
        """Get path to reference template for Pandoc, if available.

        Returns a filesystem path string if a template is configured and exists,
        otherwise returns None so we avoid passing an empty argument to Pandoc.
        """
        # For now, no default template is bundled.
        # Future enhancement: read from config and return a real path if configured.
        return None
    
    def _enhance_with_metadata(self, base_docx_path: Path, metadata: WordMetadata) -> DocxDocument:
        """Enhance base DOCX with preserved metadata."""
        doc = Document(str(base_docx_path))
        
        # Set core properties
        self._set_core_properties(doc, metadata)
        
        # Apply styles
        self._apply_styles(doc, metadata)
        
        # Set page layout
        self._set_page_layout(doc, metadata)
        
        # Configure track changes
        self._configure_track_changes(doc, metadata)
        
        # Add comments (if any)
        self._add_comments(doc, metadata)
        
        return doc
    
    def _set_core_properties(self, doc: DocxDocument, metadata: WordMetadata) -> None:
        """Set document core properties."""
        core_props = doc.core_properties
        
        if metadata.title:
            core_props.title = metadata.title
        if metadata.author:
            core_props.author = metadata.author
        if metadata.subject:
            core_props.subject = metadata.subject
        if metadata.keywords:
            core_props.keywords = ';'.join(metadata.keywords)
        if metadata.comments:
            core_props.comments = metadata.comments
    
    def _apply_styles(self, doc: DocxDocument, metadata: WordMetadata) -> None:
        """Apply preserved styles to the document."""
        if not metadata.styles:
            return
        
        try:
            # This is a simplified implementation
            # Full implementation would recreate custom styles
            
            # Set default style if specified
            if metadata.default_style and metadata.default_style in metadata.styles:
                default_style_info = metadata.styles[metadata.default_style]
                
                # Apply to all paragraphs without explicit styles
                for paragraph in doc.paragraphs:
                    if not paragraph.style.name.startswith('Heading'):
                        self._apply_style_to_paragraph(paragraph, default_style_info)
                        
        except Exception as e:
            # Style application is optional
            pass
    
    def _apply_style_to_paragraph(self, paragraph, style_info: Dict[str, Any]) -> None:
        """Apply style information to a paragraph."""
        try:
            # Apply font formatting
            if 'font' in style_info:
                font_info = style_info['font']
                
                for run in paragraph.runs:
                    if 'name' in font_info:
                        run.font.name = font_info['name']
                    if 'size' in font_info:
                        run.font.size = Pt(font_info['size'])
                    if 'bold' in font_info:
                        run.font.bold = font_info['bold']
                    if 'italic' in font_info:
                        run.font.italic = font_info['italic']
            
            # Apply paragraph formatting
            if 'paragraph_format' in style_info:
                para_info = style_info['paragraph_format']
                pf = paragraph.paragraph_format
                
                if 'alignment' in para_info:
                    # Map alignment string to enum
                    alignment_map = {
                        'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                        'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
                        'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
                        'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
                    }
                    alignment = alignment_map.get(para_info['alignment'].upper())
                    if alignment:
                        pf.alignment = alignment
                
                if 'left_indent' in para_info:
                    pf.left_indent = Inches(para_info['left_indent'] / 72)  # Convert from points
                if 'right_indent' in para_info:
                    pf.right_indent = Inches(para_info['right_indent'] / 72)
                if 'first_line_indent' in para_info:
                    pf.first_line_indent = Inches(para_info['first_line_indent'] / 72)
                if 'line_spacing' in para_info:
                    pf.line_spacing = para_info['line_spacing']
                    
        except Exception:
            # Individual style application failures are non-critical
            pass
    
    def _set_page_layout(self, doc: DocxDocument, metadata: WordMetadata) -> None:
        """Set page layout from preserved settings."""
        if not doc.sections:
            return
        
        section = doc.sections[0]  # Use first section
        
        try:
            # Set page margins
            if metadata.page_margins:
                margins = metadata.page_margins
                if 'top' in margins:
                    section.top_margin = Inches(margins['top'])
                if 'bottom' in margins:
                    section.bottom_margin = Inches(margins['bottom'])
                if 'left' in margins:
                    section.left_margin = Inches(margins['left'])
                if 'right' in margins:
                    section.right_margin = Inches(margins['right'])
                if 'header_distance' in margins:
                    section.header_distance = Inches(margins['header_distance'])
                if 'footer_distance' in margins:
                    section.footer_distance = Inches(margins['footer_distance'])
            
            # Set page size
            if metadata.page_size:
                size = metadata.page_size
                if 'width' in size:
                    section.page_width = Inches(size['width'])
                if 'height' in size:
                    section.page_height = Inches(size['height'])
                    
        except Exception:
            # Page layout errors are non-critical
            pass
    
    def _configure_track_changes(self, doc: DocxDocument, metadata: WordMetadata) -> None:
        """Configure track changes setting."""
        if not metadata.track_changes_enabled:
            return
        
        try:
            # Access document part to set track changes
            # This requires low-level XML manipulation
            doc_part = doc.part
            settings = doc_part._element.xpath('//w:settings', namespaces=self.namespaces)
            
            if settings:
                settings_elem = settings[0]
                
                # Create track changes element if it doesn't exist
                track_rev = settings_elem.xpath('w:trackRevisions', namespaces=self.namespaces)
                if not track_rev:
                    track_rev_elem = ET.Element(f'{{{self.namespaces["w"]}}}trackRevisions')
                    settings_elem.append(track_rev_elem)
                    
        except Exception:
            # Track changes configuration is optional
            pass
    
    def _add_comments(self, doc: DocxDocument, metadata: WordMetadata) -> None:
        """Add preserved comments back to the document."""
        if not metadata.document_comments:
            return
        
        # This is a complex operation that requires:
        # 1. Creating comments part
        # 2. Adding comment elements
        # 3. Linking comments to text ranges
        # For now, we'll skip this as it's quite involved
        pass
    
    def _inject_xml_fragments(
        self, 
        doc: DocxDocument, 
        fragments: XMLFragments,
        output_path: Path
    ) -> Path:
        """Re-inject preserved XML fragments into the DOCX."""
        # Save the enhanced document first
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_docx.close()
        temp_path = Path(temp_docx.name)
        
        doc.save(str(temp_path))
        
        # Now modify the DOCX ZIP to inject fragments
        if self._has_fragments_to_inject(fragments):
            final_path = self._inject_fragments_into_zip(temp_path, fragments, output_path)
        else:
            final_path = temp_path
        
        return final_path
    
    def _has_fragments_to_inject(self, fragments: XMLFragments) -> bool:
        """Check if there are any fragments to inject."""
        return bool(
            fragments.headers_footers or
            fragments.footnotes or
            fragments.endnotes or
            fragments.complex_elements or
            fragments.embedded_objects
        )
    
    def _inject_fragments_into_zip(
        self, 
        source_path: Path, 
        fragments: XMLFragments,
        output_path: Path
    ) -> Path:
        """Inject XML fragments back into the DOCX ZIP."""
        # Create output ZIP
        with zipfile.ZipFile(source_path, 'r') as source_zip:
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as output_zip:
                
                # Copy all existing files
                for item in source_zip.infolist():
                    data = source_zip.read(item.filename)
                    output_zip.writestr(item, data)
                
                # Inject headers and footers
                self._inject_headers_footers(output_zip, fragments)
                
                # Inject footnotes and endnotes
                self._inject_notes(output_zip, fragments)
                
                # Inject embedded objects
                self._inject_embedded_objects(output_zip, fragments)
                
                # Update relationships if needed
                self._update_relationships(output_zip, fragments)
        
        return output_path
    
    def _inject_headers_footers(self, zip_file: zipfile.ZipFile, fragments: XMLFragments) -> None:
        """Inject header and footer XML fragments."""
        for fragment_id, xml_content in fragments.headers_footers.items():
            # Determine filename based on fragment ID
            if fragment_id.startswith('header'):
                filename = f"word/{fragment_id}.xml"
            elif fragment_id.startswith('footer'):
                filename = f"word/{fragment_id}.xml"
            else:
                continue
            
            # Write the XML content
            zip_file.writestr(filename, xml_content.encode('utf-8'))
    
    def _inject_notes(self, zip_file: zipfile.ZipFile, fragments: XMLFragments) -> None:
        """Inject footnotes and endnotes."""
        # Collect all footnotes
        if fragments.footnotes:
            footnotes_xml = self._create_footnotes_xml(fragments.footnotes)
            zip_file.writestr("word/footnotes.xml", footnotes_xml.encode('utf-8'))
        
        # Collect all endnotes
        if fragments.endnotes:
            endnotes_xml = self._create_endnotes_xml(fragments.endnotes)
            zip_file.writestr("word/endnotes.xml", endnotes_xml.encode('utf-8'))
    
    def _create_footnotes_xml(self, footnotes: Dict[str, str]) -> str:
        """Create footnotes.xml from individual footnote fragments."""
        # This is a simplified implementation
        # Real implementation would need to properly reconstruct the footnotes XML structure
        
        root = ET.Element(f'{{{self.namespaces["w"]}}}footnotes')
        root.set('xmlns:w', self.namespaces['w'])
        
        for note_id, note_xml in footnotes.items():
            try:
                # Parse and append each footnote
                note_elem = ET.fromstring(note_xml)
                root.append(note_elem)
            except ET.ParseError:
                # Skip malformed footnotes
                continue
        
        return ET.tostring(root, encoding='unicode')
    
    def _create_endnotes_xml(self, endnotes: Dict[str, str]) -> str:
        """Create endnotes.xml from individual endnote fragments."""
        root = ET.Element(f'{{{self.namespaces["w"]}}}endnotes')
        root.set('xmlns:w', self.namespaces['w'])
        
        for note_id, note_xml in endnotes.items():
            try:
                note_elem = ET.fromstring(note_xml)
                root.append(note_elem)
            except ET.ParseError:
                continue
        
        return ET.tostring(root, encoding='unicode')
    
    def _inject_embedded_objects(self, zip_file: zipfile.ZipFile, fragments: XMLFragments) -> None:
        """Inject embedded objects and media."""
        for object_id, object_data in fragments.embedded_objects.items():
            # Determine appropriate path based on object type
            if object_id.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                filename = f"word/media/{object_id}"
            else:
                filename = f"word/embeddings/{object_id}"
            
            zip_file.writestr(filename, object_data)
    
    def _update_relationships(self, zip_file: zipfile.ZipFile, fragments: XMLFragments) -> None:
        """Update relationship files if needed for injected fragments."""
        # This is a complex operation that would involve:
        # 1. Reading existing relationship files
        # 2. Adding relationships for new fragments
        # 3. Updating the files in the ZIP
        # For now, we'll skip this as it requires detailed relationship management
        pass
    
    def validate_output(self, output_path: Path, original_model: DocumentModel) -> Dict[str, Any]:
        """
        Validate the generated DOCX file.
        
        Returns:
            Dictionary with validation results.
        """
        validation_result = {
            'file_created': output_path.exists(),
            'file_valid': False,
            'size_bytes': 0,
            'content_preserved': False,
            'metadata_preserved': False,
            'issues': [],
        }
        
        try:
            if not output_path.exists():
                validation_result['issues'].append("Output file not created")
                return validation_result
            
            validation_result['size_bytes'] = output_path.stat().st_size
            
            # Try to open the file to verify it's valid
            try:
                test_doc = Document(str(output_path))
                validation_result['file_valid'] = True
                
                # Check if content is preserved
                if test_doc.paragraphs:
                    validation_result['content_preserved'] = True
                else:
                    validation_result['issues'].append("No paragraphs found in output")
                
                # Check if metadata is preserved
                core_props = test_doc.core_properties
                if (core_props.title == original_model.word_metadata.title or
                    core_props.author == original_model.word_metadata.author):
                    validation_result['metadata_preserved'] = True
                else:
                    validation_result['issues'].append("Metadata not preserved")
                
            except Exception as e:
                validation_result['issues'].append(f"Invalid DOCX file: {str(e)}")
        
        except Exception as e:
            validation_result['issues'].append(f"Validation error: {str(e)}")
        
        return validation_result
